# -*- coding: utf-8 -*-
"""
Created on Wed Mar 22 17:56:22 2023

@author: Jinliang
"""
from word_reader import load_word, save_word
from excel_reader import load_excel, isNan
from docx.oxml.ns import qn
from convert_number import convert
import argparse

# 需要保留两位有效数字的列数据
round_two_list = ["应收权益1-本利和", "房屋总价2", "抵债金额（总）", "剩余购房款（不含首付）",
                  "对应首付金额（元）", "乙方1产品1剩余本金", "乙方1产品1剩余收益",
                  "乙方1产品1转让本金", "乙方1产品1转让收益", '房源建面', '剩余应付房款']


def replace_value_in_str(para, excel_dic, number):
    """
    针对一段文本，替换其中用##【】##括起来的值，替换为excel中的实际值
        1. 定位每两个##的位置
        2. 用其当key，取出excel中的value
        3. 直接原地替换para
    """
    count = 1
    index_temp = 0
    while True:
        try:
            index = para.index('##', index_temp)
            if count%2 == 0:
                # 直接在这里调换
                index_temp -= 1  # 恢复到之前的index
                
                key = para[index_temp+2: index]  # +2表示跳过开头的##
                
                if "合同签署年月1-1_年" in key:
                    value = excel_dic.get(key[:-2])[number].year  ## -2表示把_年去掉
                elif "合同签署年月1-1_月" in key:
                    value = excel_dic.get(key[:-2])[number].month  ## -2表示把_月去掉
                elif "阿拉伯数字" in key:  # 处理金额转成大写数字
                    related_key = key.split('_')[1]
                    raw_value = excel_dic.get(related_key)[number]
                    value = convert(str(raw_value))
                else:
                    value = excel_dic.get(key)[number]
                
                if key in round_two_list:
                    res = str(format(float(value), ","))
                    ## 这里要强行补充末尾的0
                    while len(res.split('.')[1]) < 2:
                        res += '0'  # 搞定
                    
                    value = str(res)
                
                para = para[:index_temp] + str(value) + para[index+2:]  # 把头尾的##都干掉了
                index_temp = 0  # 又要从0开始，因为之前的##被替换掉了
            else:
                index_temp = index+1  # 临时存下当前的index
            count += 1
        except ValueError:
            ## 当执行到所有的##都被替换完毕后，index(##)报错退出
            break
        
    return para


def extract_excel_to_word(word_path, excel_path, building):
    excel_dic = load_excel(excel_path)
    total_num = len(excel_dic.get('乙方1')) ## 合同总数，但不一定是真实的
    count = 0
    for number in range(total_num):
        # 无内容的记录，直接跳过
        name = excel_dic.get('乙方1')[number]
        if isNan(name):
            continue
        count += 1
        
        # 添加模板
        word_doc_mode = load_word(word_path)  # 这个东西是要反复添加的
        
        ## 先改paragraphs
        for para in word_doc_mode.paragraphs:
            if len(para.runs) == 0:
                continue  # 段落无内容，直接跳过
                
            # 默认一个para内，字体不变，字号不变
            run = para.runs[0]
            font_name = run.font.name
            font_size = run.font.size
            if_bold = run.font.bold
            
            para_new = replace_value_in_str(para.text, excel_dic, number)
            if para_new == para.text:
                continue
            para.text = para_new
            
            # 针对当前para的每个run，都要设置成对应的字体和字号
            for run in para.runs:
                run.font.name = font_name
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), font_name)
                
                run.font.size = font_size
                run.font.bold = if_bold
                
        ## 再改table
        for table in word_doc_mode.tables:
            for row in table.rows:
                for cell in row.cells:
                    
                    if len(cell.paragraphs) == 0 or len(cell.paragraphs[0].runs) == 0:
                        continue
                    
                    font_name = cell.paragraphs[0].runs[0].font.name
                    font_size = cell.paragraphs[0].runs[0].font.size
                    if_bold = cell.paragraphs[0].runs[0].font.bold
                    
                    cell_new = replace_value_in_str(cell.text, excel_dic, number)
                    cell.text = cell_new
                    
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = font_name
                            r = run._element.rPr.rFonts
                            r.set(qn('w:eastAsia'), font_name)
                            
                            run.font.size = font_size
                            run.font.bold = if_bold

        room_number = excel_dic.get('房号')[number]
        name = excel_dic.get('乙方1')[number]
        if building == '碧海云天':
            file_name = 'DF-00-' + building + '-' + room_number + '-抵房协议-' + name + '.docx'
        else:
            file_name = 'DF-HUIZ-0-' + building + '-' + room_number + '-抵房协议-' + name + '.docx'
        save_path = 'output/word_contract/' + file_name
        save_word(word_doc_mode, save_path)
    
    print('共生成合同文件数目：', count)
    
if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('building', type=str) # 楼盘名称：[碧海云天, 云筑]
    parser.add_argument('excel_path', type=str) # 需要提取的excel文件
    args = parser.parse_args()
    
    print("从【" + args.excel_path + "】中抽取信息，自动填入word模板中...")
    
    if args.building == '碧海云天':
        word_path = "data/word_mode/bihaiyuntian.docx"  # 统一的模板
    if args.building == '云筑':
        word_path = "data/word_mode/yunzhu.docx"  # 统一的模板
    extract_excel_to_word(word_path, args.excel_path, args.building)
    
    print("全部记录生成完毕，结果已写入output/word_contract文件夹")

"""
待解决
1. 多个乙方的   
2. 多个产品的
    要算加和，文字部分可能有点问题，会事先规定好吗，还是批量生成
"""

